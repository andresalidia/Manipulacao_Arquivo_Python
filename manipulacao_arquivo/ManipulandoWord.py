from docx import Document

# Add informações aos arquivo 
doc = Document() # Abrindo documento
doc.add_heading("Exemplo de Título", level=1)
doc.add_paragraph("Exemplo de paragrafo")



# Add Tabela
tabel = doc.add_table(rows=3,cols=3)
tabel.cell(0,0).text ="Matrícula"
tabel.cell(0,1).text = "Nome"
tabel.cell(0,2).text ="Data de Nascimento"
tabel.cell(1,0).text = "20250002"
tabel.cell(1,1).text = "Andresa Lídia"
tabel.cell(1,2).text = "10/05/2006"
tabel.cell(2,0).text = "20250003"
tabel.cell(2,1).text = "Faustino jr"
tabel.cell(2,2).text = "05/12/2004"

#criando documeto
doc.save('PrimeiroDocWordPython.docx')


#lendo texto em documento
for paragraph in doc.paragraphs:
    print(paragraph.text)

#lendo tabela
for tabela in doc.tables:
    for linha in tabela.rows:
        dados_linha = [celula.text.strip() for celula in linha.cells]
        print(dados_linha)

# Mostrar de forma formatada

from tabulate import tabulate #importando bibliotedca de formatação de tabela


dados = []
for tabela in doc.tables:
    for linhas in tabela.rows:
        dados_linha = [celula.text.strip() for celula in linhas.cells]
        dados.append(dados_linha)

# Separando o cabeçalo e o corpo da tabela
cabecalho = dados[0]
corpo = dados[1:]

# Mostrar formatado no terminal

print(tabulate(corpo,headers=cabecalho, tablefmt="github"))

#Exercícios1:
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()# iniciando documento
#Escrevendo Titulo
titulo = doc.add_heading()
titulo_run = titulo.add_run("Titulo Formatada")
#formatando titulo
titulo_run.bold = True
titulo_run.underline = True
titulo_run.font.size = Pt(18)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo.paragraph_format.space_after= Pt(20)

# Escrendo Paragrafos
paragrafo1 = doc.add_paragraph()
paragrafo2 = doc.add_paragraph()

paragrafo1_run = paragrafo1.add_run("Sit et ullamco sit ex magna irure cupidatat ex veniam pariatur exercitation ex sunt. Eu quis esse esse exercitation sunt duis aute elit mollit adipisicing dolor nisi et non. Reprehenderit ut cillum dolor Lorem do. Aliqua exercitation nisi qui non labore qui occaecat proident veniam commodo ad.")
paragrafo2_run = paragrafo2.add_run("Sit et ullamco sit ex magna irure cupidatat ex veniam pariatur exercitation ex sunt. Eu quis esse esse exercitation sunt duis aute elit mollit adipisicing dolor nisi et non. Reprehenderit ut cillum dolor Lorem do. Aliqua exercitation nisi qui non labore qui occaecat proident veniam commodo ad.")
#Formatando Paragrafo1
paragrafo1.paragraph_format.first_line_indent = Inches(0.5)# indentação
paragrafo1_run.font.size = Pt(12)# Tamanho da Fonte
paragrafo1_run.italic = True # Formatação da fonte
paragrafo1.paragraph_format.space_after= Pt(20)# Espaço depois do paragrafo

#Formatando Paragrafo2
paragrafo2.paragraph_format.first_line_indent = Inches(0.5)# indentação
paragrafo2_run.font.size = Pt(12)
paragrafo2_run.italic = True

#Salvando documento
doc.save("documento.docx")


#Exercício 2:

doc = Document("texto.docx")# Abrindo documento

# Percorendo o documento e printando
for paragrafos in doc.paragraphs:
    print(paragrafos.text)


#Exercício 3:

doc = Document("documento.docx")

total_palavras = 0

for paragrafos in doc.paragraphs:
    palavras = paragrafos.text.split()
    total_palavras+=len(palavras)

print(f"O total de palavras no Documento 'documento.docx' é: {total_palavras}")

# Exercício 4:
doc = Document()

# Adiciona uma tabela com 5 linhas e 3 colunas
tabela = doc.add_table(rows=5, cols=3)

# Preenche as células com "Linha X, Coluna Y"
for x in range(5):  # Linhas
    for y in range(3):  # Colunas
         tabela.cell(f"Linha {x}, Coluna {y}")
        

# Salva o documento
doc.save("tabela.docx")

print("Tabela criada e salva como 'tabela.docx'")

#Exercício 5:

doc = Document("documento.docx") #Abre documento

for paragrafo in doc.paragraphs:
    if "Lorem" in paragrafo.text:
        novo_paragrafo = paragrafo.text.replace("Lorem", "Python")
        # removendo o conteúdo antigo e add o novo
        paragrafo.clear()
        paragrafo.add_run(novo_paragrafo)

# Salvando o novo documento
doc.save("novodocumento.docx")